# -*- coding: utf-8 -*-
from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_para_format(paragraph, first_line=False, line=1.5, before=0, after=0):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Pt(21) if first_line else None


def add_text(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False, bold=False, before=0, after=0):
    p = doc.add_paragraph()
    set_para_format(p, first_line=first_line, before=before, after=after)
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    return p


def setup_doc(header_text, footer_code):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(header_text)
    set_run_font(hr)

    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = fp.add_run(footer_code + "        ")
    set_run_font(fr)
    fp2 = section.footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fr2 = fp2.add_run("2023.03")
    set_run_font(fr2)
    section.footer.add_paragraph()
    return doc


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "808080")


def write_claims(content, out_dir):
    doc = setup_doc("权利要求书", "100001")
    for claim in content["claims"]:
        add_text(doc, claim, first_line=False, after=4)
    path = out_dir / "100001权利要求书.docx"
    doc.save(path)
    return path


def write_description(content, out_dir):
    doc = setup_doc("说明书", "100002")
    for item in content["description"]:
        typ = item.get("type", "paragraph")
        text = item["text"]
        if typ == "title":
            add_text(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        elif typ == "heading":
            add_text(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=4)
        elif typ == "subheading":
            add_text(doc, text, first_line=False, before=4, after=2)
        else:
            add_text(doc, text, first_line=True, after=2)

    table_data = content.get("table")
    if table_data and table_data.get("rows"):
        add_text(doc, table_data.get("caption", "表1"), first_line=False, before=4, after=2)
        rows = table_data["rows"]
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        set_table_borders(table)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                set_para_format(p, line=1.2)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or c_idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(str(val))
                set_run_font(run, bold=(r_idx == 0))

    path = out_dir / "100002说明书.docx"
    doc.save(path)
    return path


def add_picture(doc, image_path, max_width_cm=16.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=4, after=4)
    p.add_run().add_picture(str(image_path), width=Cm(max_width_cm))


def write_drawings(content, figures_dir, out_dir):
    doc = setup_doc("说明书附图", "100003")
    drawings = content["drawings"]
    for idx, item in enumerate(drawings):
        image = figures_dir / item["file"]
        if not image.exists():
            raise FileNotFoundError(image)
        add_picture(doc, image)
        add_text(doc, item["caption"], align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
        if idx != len(drawings) - 1:
            doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
    path = out_dir / "100003说明书附图.docx"
    doc.save(path)
    return path


def write_abstract(content, figures_dir, out_dir):
    doc = setup_doc("说明书摘要", "100004")
    abstract = content["abstract"]
    text = abstract["text"]
    if len(text) > 300:
        raise ValueError(f"Abstract is {len(text)} Chinese characters; keep it <=300 by default.")
    add_text(doc, text, first_line=True, after=8)
    add_text(doc, "摘要附图", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    image = figures_dir / abstract["figure"]
    if not image.exists():
        raise FileNotFoundError(image)
    add_picture(doc, image, max_width_cm=15.0)
    path = out_dir / "100004说明书摘要.docx"
    doc.save(path)
    return path


def media_count(path):
    with zipfile.ZipFile(path) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--content", required=True, type=Path)
    parser.add_argument("--figures", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()

    content = json.loads(args.content.read_text(encoding="utf-8"))
    args.out.mkdir(parents=True, exist_ok=True)
    for existing in args.out.glob("10000*.docx"):
        existing.unlink()

    paths = [
        write_claims(content, args.out),
        write_description(content, args.out),
        write_drawings(content, args.figures, args.out),
        write_abstract(content, args.figures, args.out),
    ]
    for path in paths:
        print(f"{path.name}: media={media_count(path)} size={path.stat().st_size}")


if __name__ == "__main__":
    main()
